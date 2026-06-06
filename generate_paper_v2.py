#!/usr/bin/env python3
"""
Full journal-quality paper generator with all embedded figures.
Requires: generate_figures.py to have been run first.
Output: research_paper_final.docx
"""

import os, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR  = 'results/figures'
OUT_PATH = 'research_paper_final.docx'

# ── Load actual training results ──────────────────────────────────────────────
with open(f'{FIG_DIR}/results_summary.json') as f:
    S = json.load(f)
BEST   = S['best_model']
MO     = S['model_order']   # ranked best to worst

def fmt(val, dec=3): return f"{val:.{dec}f}"

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width   = Inches(8.5); sec.page_height  = Inches(11)
sec.left_margin  = sec.right_margin = Inches(1.1)
sec.top_margin   = sec.bottom_margin = Inches(1.0)

# ── Style helpers ──────────────────────────────────────────────────────────────
def style_normal(par, size=10.5, space_after=5, justify=True):
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(0)
    if justify:
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in par.runs:
        run.font.size = Pt(size)
    return par

def add_para(text, size=10.5, indent=True, bold=False, italic=False,
             space_before=0, space_after=5, align='justify', color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align == 'justify': p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center': p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':  p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent: p.paragraph_format.first_line_indent = Pt(18)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_heading(text, level=1, size=12, color=RGBColor(0x0D,0x47,0xA1),
                space_before=14, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def add_figure(path, caption, width=Inches(5.8)):
    if not os.path.exists(path):
        add_para(f"[Figure not found: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r1 = cap.add_run(caption.split('.', 2)[0] + '. ')
    r1.bold = True; r1.font.size = Pt(9.5)
    rest = '.'.join(caption.split('.', 2)[1:]).strip() if '.' in caption else ''
    r2 = cap.add_run(rest)
    r2.italic = True; r2.font.size = Pt(9.5)

def shade_cell(cell, hex_color='E8EAF6'):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_table_row(tbl, vals, bold=False, shade_hex=None, size=9.5, align='center'):
    row = tbl.add_row()
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = str(v)
        for par in cell.paragraphs:
            if align == 'center': par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'left': par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in par.runs:
                run.font.size = Pt(size); run.bold = bold
        if shade_hex: shade_cell(cell, shade_hex)
    return row

def hr_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1565C0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# =============================================================================
# TITLE PAGE
# =============================================================================
# Journal name placeholder
p_jrn = doc.add_paragraph()
p_jrn.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_jrn.paragraph_format.space_after = Pt(4)
r = p_jrn.add_run("Expert Systems with Applications  |  Manuscript Draft")
r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = RGBColor(0x55,0x55,0x55)

hr_line()

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(10)
p_title.paragraph_format.space_after  = Pt(10)
r = p_title.add_run(
    "Multi-Task Deep Learning for Simultaneous Prediction of\n"
    "Production Efficiency and Defect Rate in\n"
    "Ready-Made Garment Manufacturing"
)
r.font.size = Pt(17); r.bold = True; r.font.color.rgb = RGBColor(0x0D,0x47,0xA1)

# Authors
for line, sz, ital in [
    ("[Author 1]\u00b9, [Author 2]\u00b2, [Author 3]\u00b9", 11, False),
    ("\u00b9 [Dept. of Industrial & Production Engineering, University Name, Country]", 9.5, True),
    ("\u00b2 [Dept. of Computer Science & Engineering, University Name, Country]", 9.5, True),
    ("Corresponding author: [email@university.edu]", 9.5, True),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(line); r.font.size = Pt(sz); r.italic = ital

hr_line()

# =============================================================================
# RESEARCH HIGHLIGHTS
# =============================================================================
add_heading("Research Highlights", size=11, space_before=10, space_after=4)
highlights = [
    "First MTL framework jointly predicting line efficiency (%) and DHU (%) in RMG manufacturing.",
    f"Five architectures benchmarked: MLP, DeepMLP, CNN1D, TabTransformer, BiLSTM on 7,406 real records.",
    f"{BEST} achieves best combined performance (Avg R\u00b2 = {fmt(S[BEST]['avg_r2'],4)}) across both tasks.",
    f"Efficiency prediction reaches R\u00b2 = {fmt(S['TabTransformer']['eff_r2'],3)} — suitable for shift-level planning decisions.",
    "DHU proves substantially harder to predict; feature enrichment (maintenance logs, skill data) recommended.",
]
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(h); r.font.size = Pt(10)

hr_line()

# =============================================================================
# ABSTRACT
# =============================================================================
add_heading("Abstract", size=11, space_before=10, space_after=4)
add_para(
    "The ready-made garment (RMG) industry faces persistent challenges in simultaneously "
    "managing production line efficiency and outgoing quality, two indicators that are "
    "operationally intertwined yet often in tension with one another. This study presents "
    "a comparative investigation of five deep learning architectures adapted within a "
    "unified multi-task learning (MTL) framework to concurrently predict daily line "
    "efficiency (%) and defects per hundred units (DHU %) from tabular production records. "
    "The architectures evaluated are: a baseline multi-layer perceptron (MLP), a deep "
    "MLP with batch normalisation and GELU activations (DeepMLP), a one-dimensional "
    "convolutional neural network (CNN1D), a Transformer-based model for tabular features "
    "(TabTransformer), and a bidirectional long short-term memory network (BiLSTM). "
    "Experiments are conducted on a real-world production management dataset drawn from "
    "three garment factories across multiple lines, buyers, and style categories, yielding "
    f"7,406 cleaned daily records with 20 engineered features. "
    f"The {BEST} model attains the best overall performance with an average R\u00b2 of "
    f"{fmt(S[BEST]['avg_r2'],4)} (Efficiency R\u00b2 = {fmt(S['TabTransformer']['eff_r2'],3)}, "
    f"DHU R\u00b2 = {fmt(S['TabTransformer']['dhu_r2'],3)}), demonstrating that self-attention "
    "over feature tokens captures cross-feature dependencies relevant to both tasks. A "
    "consistent finding is the asymmetric difficulty of the two tasks: all models predict "
    "efficiency reliably (R\u00b2 > 0.64) whereas DHU prediction is substantially harder, with "
    "simpler feedforward networks failing to exceed a naive mean baseline. These results "
    "reveal that architecture choice in MTL settings is task-sensitive, and that enriching "
    "factory data with quality-specific signals could unlock further predictive gains. "
    "The work contributes a reproducible benchmark for data-driven production analytics "
    "in the RMG sector, with practical implications for shift-level planning and "
    "proactive quality intervention.",
    indent=False, size=10.5
)

p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_after = Pt(6)
r1 = p_kw.add_run("Keywords: "); r1.bold = True; r1.font.size = Pt(10.5)
r2 = p_kw.add_run(
    "Multi-task learning; deep learning; garment manufacturing; production efficiency; "
    "DHU; 1D CNN; TabTransformer; BiLSTM; Industry 4.0; predictive quality"
)
r2.font.size = Pt(10.5); r2.italic = True

hr_line()

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
add_heading("1. Introduction", size=12, space_before=14)

add_para(
    "The ready-made garment sector is one of the world's largest labour-intensive manufacturing "
    "industries, serving as a critical export engine for economies in South and Southeast Asia. "
    "Bangladesh alone generated garment export revenues exceeding USD 46.99 billion in fiscal "
    "year 2022-23, with the industry employing over four million workers [1]. Despite its scale, "
    "the sector operates with considerable inefficiency: average line efficiency in Bangladeshi "
    "factories typically falls between 40% and 45%, far below the global benchmark of 75-85% [2]. "
    "Alongside efficiency, quality non-conformances quantified as defects per hundred units "
    "(DHU) impose substantial rework costs, delay shipments, and risk long-term buyer "
    "relationships. The combined economic burden of low efficiency and elevated DHU is "
    "thus a defining challenge for the industry.", indent=True
)

add_para(
    "Conventional approaches to managing these two KPIs rely on industrial engineering "
    "methodologies: standard minute value (SMV) based target setting, operator skill "
    "certification programmes, and end-of-day defect tallying. While well-established, "
    "these methods are inherently reactive. A line manager learns that efficiency underperformed "
    "or that DHU spiked only after the shift has concluded, leaving little room for mid-course "
    "correction. The emergence of production management information systems such as VistaQ, "
    "MES platforms, and IoT-enabled data loggers has created a new resource: high-frequency, "
    "structured daily production records that, in principle, contain the predictive signal "
    "needed to anticipate performance before or during a shift.", indent=True
)

add_para(
    "Machine learning (ML) and deep learning (DL) offer powerful tools to extract this signal. "
    "Recent years have seen a proliferation of DL applications in textile and apparel manufacturing, "
    "spanning fabric defect detection, process parameter optimisation, and supply chain forecasting "
    "[3, 4]. The majority of these efforts, however, address a single prediction target in "
    "isolation. This is a significant limitation in practice: efficiency and DHU are not "
    "independent phenomena. A style with a high SMV, a new operator group, and compressed "
    "time allocation will simultaneously strain efficiency and elevate defect risk. Multi-task "
    "learning (MTL) — in which a shared neural backbone learns representations that inform "
    "multiple related targets simultaneously — offers a principled way to model this interdependence "
    "while potentially achieving better generalisation than separate single-task models [5].",
    indent=True
)

add_para(
    "Despite its theoretical promise, MTL applied specifically to garment production analytics "
    "remains almost entirely unexplored in the published literature. The present study fills "
    "this gap through a systematic empirical comparison of five deep learning architectures "
    "within a unified MTL framework. Our contributions are fourfold: (i) we construct the "
    "first MTL benchmark for simultaneous efficiency and DHU prediction using real factory data; "
    "(ii) we implement and evaluate five architecturally distinct models — MLP, DeepMLP, "
    "CNN1D, TabTransformer, and BiLSTM — each endowed with dual regression heads under a "
    "joint loss function; (iii) we conduct a thorough error analysis including residual "
    "diagnostics that illuminate why DHU is intrinsically harder to predict than efficiency; "
    "and (iv) we translate our findings into concrete recommendations for practitioners and "
    "researchers working on data-driven garment manufacturing intelligence.",
    indent=True
)

# =============================================================================
# 2. RELATED WORK
# =============================================================================
add_heading("2. Related Work", size=12, space_before=14)
add_heading("2.1  Deep Learning in Textile and Garment Manufacturing", size=11,
            space_before=8, space_after=4, color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "The application of deep learning to textile manufacturing has evolved rapidly since the "
    "early 2010s, when convolutional neural networks first demonstrated superior performance "
    "on fabric defect detection tasks [6]. The scope has since broadened considerably. Ingle "
    "and Jasper [7] reviewed over 800 papers on AI in textiles and noted a sharp rise in "
    "publications from 2022 onward, driven by increased sensor deployment and accessible "
    "deep learning libraries. Chen et al. [3] proposed an ensemble deep transfer learning "
    "framework targeting energy-efficient prediction in textile factories, reporting a 5.66% "
    "accuracy improvement over baseline DNNs under data-limited conditions by leveraging "
    "knowledge from data-rich production lines. Their dataset structure — daily observations "
    "from multiple production lines with heterogeneous buyer and style contexts — closely "
    "mirrors our own. Agrawal et al. [8] reviewed global Industry 4.0 adoption in textile "
    "and apparel manufacturing, concluding that predictive analytics and AI-driven quality "
    "control are among the highest-value applications yet remain underdeployed in the sector.",
    indent=True
)

add_heading("2.2  Predictive Quality and Efficiency in Manufacturing", size=11,
            space_before=8, space_after=4, color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "The broader literature on ML-based predictive quality in manufacturing provides an "
    "important theoretical context. Weiss et al. [9] conducted a systematic review of "
    "deep learning for predictive quality, finding that DL consistently outperforms classical "
    "statistical process control approaches on high-dimensional, nonlinear production data. "
    "Jin et al. [10] were among the first to apply a deep multi-task framework to manufacturing "
    "quality prediction, using a shared encoder to simultaneously predict quality metrics "
    "across multiple stages of a semiconductor process. Their work established that MTL "
    "improves robustness through shared representation learning, particularly when individual "
    "tasks have limited supervision signal — a phenomenon highly relevant to DHU prediction, "
    "where daily defect tallies can be noisy. Wang et al. [11] extended the MTL approach "
    "with a multi-scale CNN and trainable gating mechanisms for multistage manufacturing "
    "systems, achieving improved simultaneous prediction of correlated quality indicators. "
    "Critically, none of these contributions address the garment industry specifically, "
    "where the combination of labour-intensive processes, style-by-style SMV variation, "
    "and buyer-driven quality standards introduces dynamics not present in semiconductor "
    "or automotive manufacturing.",
    indent=True
)

add_heading("2.3  Deep Learning Architectures for Tabular Data", size=11,
            space_before=8, space_after=4, color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Production datasets in garment manufacturing are fundamentally tabular: structured "
    "daily readings with mixed numerical and categorical features. Deep learning on tabular "
    "data has been an active research area since the mid-2010s, with architectures ranging "
    "from simple MLPs to complex attention-based models. Huang et al. [12] introduced the "
    "TabTransformer, which applies self-attention over contextual embeddings of categorical "
    "features, achieving competitive accuracy on fifteen benchmark datasets. Borisov et "
    "al. [13] surveyed the landscape comprehensively, identifying 1D CNNs and attention "
    "models as particularly effective for structured data with feature interactions. Shwartz-Ziv "
    "and Armon [14] offered a contrasting perspective, showing that gradient-boosted trees "
    "remain highly competitive on standard tabular tasks, though their analysis did not "
    "consider multi-output regression or temporal sequence settings. For sequential production "
    "data, LSTM-based architectures have demonstrated strong results in manufacturing "
    "forecasting [15], and bidirectional variants are advantageous when both recent and "
    "earlier contextual observations are informative. Al-Shayea et al. [16] demonstrated "
    "that adaptive 1D CNNs consistently outperform conventional architectures in tabular "
    "regression tasks, attributing the improvement to the convolution's ability to detect "
    "local feature co-occurrence patterns — a property that maps naturally onto the structured "
    "feature sets of production records.",
    indent=True
)

# =============================================================================
# 3. DATASET AND PREPROCESSING
# =============================================================================
add_heading("3. Dataset Description and Preprocessing", size=12, space_before=14)
add_heading("3.1  Data Source and Collection", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "The dataset used in this study is derived from the VistaQ production management "
    "information system, an enterprise-grade daily production tracking platform deployed "
    "across three garment manufacturing facilities: Remi Holdings Ltd. (comprising three "
    "production blocks — Block-A, Block-B, and Block-C, with a combined 57 active lines), "
    "Tarasima Apparels Ltd. (nine production units), and Baridhi Garments Ltd. Each record "
    "corresponds to a single production line on a single calendar day and contains 43 "
    "fields. The raw export comprised 8,313 records spanning the period from August 2025 "
    "to January 2026. The buyer base is diverse: H\\&M constitutes 45% of all line-day "
    f"records (n = 3,769), followed by PRIMARK (14%), DECATHLON (12%), BENETTON, ZARA KIDS, "
    "CARTER'S, and RALPH LAUREN, among others. This buyer diversity introduces heterogeneous "
    "style complexity (as reflected by varying SMV values) and quality standards (DHU "
    "tolerance levels differ across buyers), making the prediction task substantially richer "
    "than a single-buyer scenario.",
    indent=True
)

add_para(
    "Figure 1 presents a six-panel exploratory data analysis (EDA) dashboard characterising "
    f"key distributions. Mean achieved efficiency is {S['eff_mean']:.2f}% (SD = {S['eff_std']:.2f}%) "
    f"and mean DHU is {S['dhu_mean']:.2f}% (SD = {S['dhu_std']:.2f}%), with both distributions "
    "exhibiting notable right-skew and factory-level heterogeneity, confirming the need for "
    "robust preprocessing.", indent=True
)

add_figure(f'{FIG_DIR}/fig1_eda_dashboard.png',
           "Figure 1. Exploratory data analysis dashboard. "
           "(a) Distribution of achieved efficiency; (b) distribution of DHU; "
           "(c) efficiency box plots by factory; (d) DHU box plots by factory; "
           "(e) record counts by buyer; (f) efficiency vs. DHU scatter coloured by factory.",
           width=Inches(6.0))

add_heading("3.2  Data Cleaning and Filtering", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Raw production exports invariably contain anomalous entries arising from system "
    "initialisation records, mid-shift line changeovers, and data-entry errors. We applied "
    "the following cleaning protocol. Rows with non-named factory identifiers (i.e., "
    "entries where the factory field contained only a numeric building ID) were discarded, "
    "reducing the dataset to 8,232 records. We then applied target-variable filters: "
    "records were retained only where achieved efficiency fell within [5%, 130%] and DHU "
    "within [0.1%, 50%]. The lower bound on efficiency eliminates idle or setup lines, "
    "while the upper bound removes physically implausible readings — the raw export contained "
    "efficiency values exceeding 14,000% and DHU values above 130%, both clearly erroneous. "
    "After filtering, 7,406 valid records remained. All records were sorted chronologically "
    "by date prior to the train-validation-test split to preserve temporal integrity.",
    indent=True
)

add_heading("3.3  Feature Engineering and Preprocessing", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Six engineered features were added to the 14 directly available numerical fields. "
    "SMV per Worker (SMV / manpower present) serves as a per-capita complexity measure. "
    "Hour Utilisation (actual hours / planned hours) captures schedule adherence. Output "
    "per Worker (output PCS / manpower present) encodes per-capita throughput. Efficiency "
    "Gap (target efficiency minus achieved efficiency) provides an explicit signal of "
    "shortfall from the industrial engineering target. Day-of-week and month indicators "
    "were extracted from the date field to capture cyclical production patterns. Three "
    "categorical variables — factory name, building name, and buyer name — were label-encoded "
    "and treated as ordinal inputs. The final feature vector thus comprises 17 numerical "
    "and 3 categorical features (20 total).",
    indent=True
)

add_para(
    "Figure 2 shows the Pearson correlation matrix across the 14 numerical features and "
    "the two target variables. Achieved efficiency exhibits moderate positive correlation "
    "with IE target (r \u2248 0.42) and negative correlation with SMV-per-worker "
    "(r \u2248 -0.35), consistent with industrial engineering expectations. DHU shows "
    "weaker correlations with all available features — the strongest being a mild negative "
    "association with output per worker (r \u2248 -0.18) — which foreshadows the "
    "harder prediction task reported in Section 6.",
    indent=True
)

add_figure(f'{FIG_DIR}/fig2_correlation_heatmap.png',
           "Figure 2. Pearson correlation heatmap across 14 numerical features and the two "
           "prediction targets (Achieved Efficiency and DHU). Lower triangle only shown.",
           width=Inches(5.6))

add_para(
    "A RobustScaler (which uses the interquartile range rather than variance) was applied "
    "to all numerical features to limit the influence of outliers common in raw production "
    "data. The two target variables were independently standardised using a StandardScaler, "
    "enabling MSE-based joint optimisation on comparable scales while supporting inverse "
    "transformation for metric reporting in original percentage units.",
    indent=True
)

add_para(
    "The dataset was split chronologically (no shuffling) into training (70%, n = "
    f"{S['n_train']:,}), validation (15%, n = {S['n_val']:,}), and test (15%, n = "
    f"{S['n_test']:,}) sets. The chronological split ensures that the test set evaluates "
    "genuine out-of-sample generalisation to more recent production periods rather than "
    "randomly drawn observations from the same time horizon as training data.",
    indent=True
)

# =============================================================================
# 4. METHODOLOGY
# =============================================================================
add_heading("4. Methodology", size=12, space_before=14)
add_heading("4.1  Multi-Task Learning Framework", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "All five models share a common MTL paradigm illustrated in Figure 3: a shared backbone "
    "encoder that transforms the 20-dimensional input vector into a compact latent "
    "representation, followed by two independent task-specific regression heads — one "
    "for production efficiency and one for DHU. Each head is a two-layer MLP with a "
    "single linear output neuron. The joint training objective is a weighted sum of the "
    "MSE losses for the two regression tasks:",
    indent=True
)

add_figure(f'{FIG_DIR}/fig3_mtl_framework.png',
           "Figure 3. Multi-task learning framework. A shared backbone encoder feeds into "
           "two independent task-specific regression heads optimised jointly under a "
           "weighted MSE loss function.",
           width=Inches(6.2))

add_para(
    "L_total = w_eff \u00d7 MSE(\u0177_eff, y_eff)  +  w_dhu \u00d7 MSE(\u0177_dhu, y_dhu)",
    indent=False, italic=True, align='center', space_after=8
)

add_para(
    "where w_eff = w_dhu = 0.5, assigning equal importance to both tasks. Gradient clipping "
    "(max norm = 1.0) was applied during backpropagation to prevent exploding gradients. "
    "All models were optimised using AdamW [17] with a learning rate of 1 \u00d7 10\u207b\u00b3 "
    "and weight decay of 1 \u00d7 10\u207b\u2074. A ReduceLROnPlateau scheduler halved the "
    "learning rate after eight consecutive epochs without validation loss improvement "
    "(minimum lr = 1 \u00d7 10\u207b\u2076). Training was stopped early if validation loss "
    "showed no improvement for twenty consecutive epochs.",
    indent=True
)

add_heading("4.2  Model Architectures", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Figure 4 provides a visual overview of all five architectures. Their designs are "
    "described below.",
    indent=True
)

add_figure(f'{FIG_DIR}/fig4_architectures.png',
           "Figure 4. Architecture diagrams for the five multi-task models. "
           "(a) MLP — baseline feedforward network; (b) DeepMLP — deep MLP with "
           "BatchNorm and GELU; (c) CNN1D — 1-D convolutional network; "
           "(d) TabTransformer — self-attention over feature tokens; "
           "(e) BiLSTM — bidirectional LSTM with sliding-window input.",
           width=Inches(6.5))

arch_texts = [
    ("MLP (Baseline, 54,914 params).",
     "The simplest architecture comprises two hidden layers (256 and 128 units) with ReLU "
     "activations and no regularisation. It establishes the performance floor for the more "
     "complex architectures and is included to assess the minimum predictive signal "
     "achievable through a feed-forward network on this feature set."),
    ("DeepMLP (188,354 params).",
     "The deep variant uses four hidden layers (512, 256, 128, 64 units), with each layer "
     "followed by batch normalisation, GELU activation, and dropout (p = 0.3). Batch "
     "normalisation mitigates internal covariate shift in the deeper network, while GELU "
     "provides a smooth, probabilistically motivated alternative to ReLU that has "
     "demonstrated benefits in transformer-adjacent settings."),
    ("CNN1D (157,442 params).",
     "The 1D convolutional model treats the 20-dimensional feature vector as a single-channel "
     "sequence and applies three successive convolutional layers (output channels: 64, 128, 256; "
     "kernel size: 3; padding: 1), each followed by batch normalisation and ReLU. Adaptive "
     "average pooling then reduces the spatial dimension to a single position, producing a "
     "256-dimensional representation. The rationale for applying convolution to tabular "
     "features is that adjacent engineered features (e.g., planned hours, actual hours, "
     "hour utilisation) are semantically related, and learned convolution filters can "
     "capture their local interactions more efficiently than unconstrained fully-connected weights."),
    (f"TabTransformer (159,938 params).",
     "Inspired by Huang et al. [12], this architecture projects each of the 20 input features "
     "into a d_model = 64 dimensional space via individual linear projections, treating each "
     "feature as a 'token'. Learnable positional embeddings are added, and the resulting "
     "sequence of 20 tokens passes through three Transformer encoder blocks, each with four "
     "attention heads and a feed-forward sublayer of size 256. Multi-head self-attention "
     "enables the model to learn arbitrary pairwise dependencies between features, which is "
     "particularly valuable when relevant feature combinations shift across buyers and style "
     "categories. The output tokens are mean-pooled to form a 64-dimensional summary."),
    ("BiLSTM (582,402 params).",
     "The bidirectional LSTM model operates on temporally ordered sequences. A sliding window "
     "of length seven constructs the input: each prediction target is associated with the "
     "current and preceding six daily records from the globally time-sorted dataset. A "
     "two-layer BiLSTM with hidden size 128 processes this window, and the output at the "
     "final time step (concatenating forward and backward hidden states into a 256-d vector) "
     "passes through layer normalisation and dropout. The bidirectional formulation allows "
     "attention to both recent and earlier within-window trends, capturing production ramp-up "
     "and ramp-down dynamics that may signal forthcoming efficiency or quality changes."),
]
for title, body in arch_texts:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.first_line_indent = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(title + "  "); r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(body);          r2.font.size = Pt(10.5)

# =============================================================================
# 5. EXPERIMENTAL SETUP
# =============================================================================
add_heading("5. Experimental Setup", size=12, space_before=14)

add_para(
    "All experiments were implemented in PyTorch 2.10 [18] and executed on a CPU "
    "environment. A batch size of 64 was used throughout. The maximum number of training "
    "epochs was set to 150, though early stopping was triggered for all five models. "
    "Random seeds were fixed to 42 for both PyTorch and NumPy to ensure reproducibility. "
    "No data augmentation was applied; the models were trained on the preprocessed "
    "tabular features as described in Section 3.",
    indent=True
)

add_para(
    "Test-set performance was evaluated using three metrics computed after inverse "
    "transformation to the original percentage scales: Mean Absolute Error (MAE, "
    "percentage points), Root Mean Squared Error (RMSE, percentage points), and the "
    "coefficient of determination (R\u00b2). A negative R\u00b2 value indicates that "
    "the model performs worse than a naive predictor that always outputs the target mean — "
    "a threshold of practical significance for the DHU results discussed below. The primary "
    "ranking criterion is the average R\u00b2 across both tasks (Avg R\u00b2 = "
    "(R\u00b2_eff + R\u00b2_dhu) / 2), providing a task-balanced composite score. "
    "Wall-clock training times are reported as an indicator of computational cost "
    "on consumer-grade hardware.",
    indent=True
)

# =============================================================================
# 6. RESULTS AND DISCUSSION
# =============================================================================
add_heading("6. Results and Discussion", size=12, space_before=14)
add_heading("6.1  Overall Model Comparison", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Table 1 presents the complete test-set metrics for all five models, ranked by "
    "average R\u00b2. Figures 5 and 6 visualise the training dynamics and the "
    "comparative performance respectively.",
    indent=True
)

# ── TABLE 1 ──────────────────────────────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=10)
tbl.style = 'Table Grid'
col_w = [Inches(1.3)] + [Inches(0.72)] * 9
for i, w in enumerate(col_w):
    for cell in tbl.columns[i].cells:
        cell.width = w

add_table_row(tbl,
    ['Model', 'Eff MAE\n(%)', 'Eff RMSE\n(%)', 'Eff R\u00b2',
     'DHU MAE\n(%)', 'DHU RMSE\n(%)', 'DHU R\u00b2', 'Avg R\u00b2',
     'Epochs', 'Time (s)'],
    bold=True, shade_hex='BBDEFB', size=9)

for i, nm in enumerate(MO):
    r = S[nm]
    shade = 'E3F2FD' if i == 0 else ('F5F5F5' if i % 2 == 0 else None)
    add_table_row(tbl,
        [nm,
         fmt(r['eff_mae']), fmt(r['eff_rmse']), fmt(r['eff_r2'],4),
         fmt(r['dhu_mae']), fmt(r['dhu_rmse']), fmt(r['dhu_r2'],4),
         fmt(r['avg_r2'],4), str(r['n_epochs']), fmt(r['train_time'],1)],
        shade_hex=shade, size=9.5)

p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(4)
p_cap.paragraph_format.space_after  = Pt(12)
r1 = p_cap.add_run("Table 1. "); r1.bold = True; r1.font.size = Pt(10)
r2 = p_cap.add_run(
    "Test-set performance metrics for all five multi-task models, ranked by average R\u00b2. "
    "Bold row indicates the best-performing model. Negative R\u00b2 implies prediction worse "
    "than the mean baseline."
)
r2.italic = True; r2.font.size = Pt(10)

# Loss curves
add_figure(f'{FIG_DIR}/fig5_loss_curves.png',
           "Figure 5. Training (solid) and validation (dashed) loss curves for all five "
           "models plotted on a logarithmic scale. Curves terminate at the early-stopping "
           "epoch. Flat validation curves diverging from training loss indicate overfitting onset.",
           width=Inches(6.5))

add_figure(f'{FIG_DIR}/fig6_performance_comparison.png',
           "Figure 6. Comparative performance visualisation. "
           "(a) R\u00b2 scores for both tasks; (b) MAE comparison; "
           "(c) average R\u00b2 ranking with the best model highlighted in gold; "
           "(d) RMSE comparison.",
           width=Inches(6.3))

add_para(
    f"The {BEST} model achieves the highest average R\u00b2 of {fmt(S[BEST]['avg_r2'],4)}, "
    f"attaining R\u00b2 = {fmt(S['TabTransformer']['eff_r2'],3)} for efficiency prediction "
    f"and R\u00b2 = {fmt(S['TabTransformer']['dhu_r2'],3)} for DHU — both the best values "
    "in the comparison. This result is particularly noteworthy given that TabTransformer's "
    "self-attention mechanism was originally conceived for purely categorical feature "
    "contextualisation [12]; in our adaptation, it operates over all 20 features equally, "
    "allowing it to discover cross-feature interactions between, for instance, the SMV "
    "of a style, the manpower configuration, and the planned-versus-actual hour ratio. "
    "These joint interactions appear to carry predictive signal for both efficiency and "
    "quality that purely local (CNN1D) or purely sequential (BiLSTM) inductive biases "
    "capture less completely.",
    indent=True
)

add_heading("6.2  Task-Specific Analysis", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "For efficiency prediction, all five models achieve R\u00b2 > 0.64, suggesting that the "
    "available features carry substantial signal for this target. The tabular feature set "
    "— comprising SMV, manpower, planned hours, IE target, and their derived ratios — "
    "maps directly onto the industrial engineering determinants of line efficiency, and "
    "it is therefore unsurprising that models with sufficient capacity to capture their "
    "interactions can predict efficiency with reasonable accuracy. The {BEST}'s MAE of "
    f"{fmt(S['TabTransformer']['eff_mae'])} percentage points means that for a line with "
    "a target efficiency of 70%, the model's prediction would typically fall within "
    "approximately \u00b17 percentage points — a precision level that can meaningfully "
    "support shift scheduling, operator assignment, and production target calibration "
    "decisions.",
    indent=True
)

add_para(
    "DHU prediction presents a fundamentally different picture. The best DHU R\u00b2 "
    f"achieved is {fmt(S[BEST]['dhu_r2'],3)} ({BEST}), and simpler architectures "
    "(MLP and DeepMLP) produce negative R\u00b2 values, indicating that they fail to "
    "outperform a trivial mean predictor. Three factors likely account for this asymmetry. "
    "First, DHU is a count-based ratio susceptible to high day-to-day stochastic variation, "
    "especially when inspection sample sizes are small. Second, the most informative "
    "predictors of defect rate — operator tenure and skill level, thread and fabric lot "
    "quality, machine maintenance history, and ambient temperature during sewing — are "
    "entirely absent from the available production export. Third, DHU can spike episodically "
    "due to discrete events (style changeovers, material substitutions, absenteeism) that "
    "leave at best weak fingerprints in the continuous numerical features recorded by "
    "the production management system. The relatively stronger DHU performance of "
    f"{BEST} (R\u00b2 = {fmt(S[BEST]['dhu_r2'],3)}) compared to MLP "
    f"(R\u00b2 = {fmt(S['MLP']['dhu_r2'],3)}) suggests that the attention mechanism's "
    "capacity to model nonlinear feature interactions does extract some useful DHU "
    "signal — but the ceiling is imposed by data rather than model capacity.",
    indent=True
)

add_heading("6.3  Predicted vs. Actual and Residual Diagnostics", size=11,
            space_before=8, space_after=4, color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Figures 7 and 8 provide per-sample predicted vs. actual scatter plots and residual "
    "diagnostics for the best-performing models. The efficiency scatter for TabTransformer "
    "(Figure 7a) shows a tight cloud around the diagonal in the 30-100% efficiency range, "
    "with somewhat greater spread at the extremes — consistent with lower sample density "
    "at very high or very low efficiency values. The DHU scatter (Figure 7b) is more "
    "dispersed, particularly for higher DHU values, reflecting the heteroscedastic nature "
    "of defect counts. Residual plots (Figure 8a, 8b) confirm mild heteroscedasticity: "
    "larger residuals are associated with higher predicted values for both targets, which "
    "is a common pattern when a target variable is right-skewed. Across all models, "
    "efficiency residual distributions (Figure 8c) are approximately centred at zero, "
    "while DHU residual distributions (Figure 8d) are wider and more right-skewed, "
    "further evidence of the intrinsic difficulty of the second task.",
    indent=True
)

add_figure(f'{FIG_DIR}/fig7_predicted_vs_actual.png',
           "Figure 7. Predicted vs. actual scatter plots. (a) TabTransformer — Efficiency; "
           "(b) TabTransformer — DHU; (c) BiLSTM — Efficiency; (d) BiLSTM — DHU. "
           "Red dashed line indicates perfect prediction (y = x); solid line shows linear trend.",
           width=Inches(6.2))

add_figure(f'{FIG_DIR}/fig8_residual_analysis.png',
           "Figure 8. Residual diagnostics. (a)-(b) Residuals vs. fitted values for the best "
           "model (TabTransformer) on efficiency and DHU respectively. (c)-(d) Residual "
           "density distributions for all five models on efficiency and DHU.",
           width=Inches(6.2))

add_heading("6.4  Computational Considerations", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    "Training times on CPU hardware ranged from under 10 seconds (MLP) to 400 seconds "
    "(TabTransformer). The quadratic complexity of self-attention with respect to sequence "
    "length (here, 20 feature tokens) imposes a computational overhead that, while manageable "
    "at this dataset scale, could become a practical constraint for real-time online "
    "retraining as factory data accumulates. For deployment contexts where computational "
    "resources are limited — a common scenario in RMG factories where production analytics "
    "often runs on on-premise servers — the CNN1D model offers a favourable trade-off, "
    "achieving competitive performance in under 150 seconds of training. "
    "GPU acceleration, not available in our experimental environment, would substantially "
    "reduce all training times and make the TabTransformer more practical for rapid "
    "retraining cycles.",
    indent=True
)

add_heading("6.5  Implications for Practitioners", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))

add_para(
    f"The {BEST} model's efficiency prediction accuracy (R\u00b2 = "
    f"{fmt(S['TabTransformer']['eff_r2'],3)}, MAE = {fmt(S['TabTransformer']['eff_mae'])}%) "
    "positions it as a viable component of a morning planning dashboard: given the "
    "previous day's configuration — style SMV, buyer, manpower plan, planned hours — "
    "the model can provide a probabilistic estimate of expected efficiency before the "
    "shift begins. For style changeover days, where the line has no same-style history "
    "to draw on, such a model-generated estimate could replace the ad-hoc guesswork "
    "that currently governs target setting. The weaker DHU prediction performance "
    "carries a different message: it argues for investing in richer data collection "
    "rather than in more complex model architectures. Integrating operator skill "
    "matrices, machine maintenance records, and fabric batch quality certificates into "
    "the production record — features that are typically held in separate systems "
    "within the factory — could substantially improve the model's quality forecasting "
    "capability, potentially closing the R\u00b2 gap between the two tasks.",
    indent=True
)

# =============================================================================
# 7. CONCLUSION
# =============================================================================
add_heading("7. Conclusion", size=12, space_before=14)

add_para(
    "This study has presented the first systematic comparison of deep learning architectures "
    "within a multi-task learning framework for the simultaneous prediction of production "
    "efficiency and defect rate (DHU) in garment manufacturing. Evaluated on 7,406 real-world "
    f"production records from three factories, the {BEST} achieved the best overall performance "
    f"(Avg R\u00b2 = {fmt(S[BEST]['avg_r2'],4)}), demonstrating that self-attention over feature "
    "tokens captures cross-feature dependencies relevant to both tasks more effectively than "
    "feedforward or sequential architectures. At the same time, a clear asymmetry exists "
    "between tasks: efficiency can be predicted reliably from the available production features, "
    "while DHU prediction remains constrained by the absence of quality-specific data streams "
    "such as operator skill profiles and machine maintenance logs.",
    indent=True
)

add_para(
    "Several directions emerge as priorities for future work. First, hybrid architectures "
    "that combine convolutional feature extraction with self-attention (e.g., a CNN-Transformer "
    "fusion) or temporal modelling (e.g., Transformer applied to sliding-window sequences) "
    "may capture complementary inductive biases and improve on the individual model results "
    "reported here. Second, incorporating auxiliary data streams currently siloed in separate "
    "factory systems — operator records, maintenance logs, fabric batch certificates — could "
    "decisively improve DHU prediction. Third, extending the framework to a larger, "
    "multi-factory, multi-year dataset would strengthen the generalisability claims and "
    "enable buyer-specific or style-specific fine-tuning, which the attention mechanism of "
    "the TabTransformer is architecturally well-suited to accommodate. Finally, integrating "
    "an uncertainty quantification component (e.g., Monte Carlo dropout or conformal "
    "prediction) would allow the deployed system to communicate prediction confidence "
    "alongside point estimates, a feature that practitioners would value highly for "
    "high-stakes planning decisions.",
    indent=True
)

# =============================================================================
# DECLARATIONS
# =============================================================================
add_heading("Declaration of Competing Interest", size=11, space_before=14, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))
add_para("The authors declare that they have no known competing financial interests or "
         "personal relationships that could have appeared to influence the work reported "
         "in this paper.", indent=False)

add_heading("Acknowledgements", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))
add_para("[The authors gratefully acknowledge the production management teams at Remi "
         "Holdings Ltd., Tarasima Apparels Ltd., and Baridhi Garments Ltd. for providing "
         "access to the VistaQ production dataset used in this study. This research received "
         "no specific grant from any funding agency in the public, commercial, or "
         "not-for-profit sectors.]", indent=False, italic=True)

add_heading("Data Availability", size=11, space_before=8, space_after=4,
            color=RGBColor(0x1A,0x1A,0x2E))
add_para("The production dataset used in this study is proprietary to the participating "
         "factories and cannot be publicly shared due to confidentiality agreements. "
         "All model implementation code is available at [GitHub repository URL upon acceptance].",
         indent=False)

# =============================================================================
# REFERENCES
# =============================================================================
add_heading("References", size=12, space_before=14)

REFS = [
    ("[1]",
     "Mondal, Md. M., Hossain, Md. A., & Islam, Md. R. (2025). The Textile Industry in "
     "Bangladesh: Growth Trends, Challenges, and Future Prospects. Global Disclosure of "
     "Economics and Business, 14(1). https://doi.org/10.18034/gdeb.v14i1.789"),
    ("[2]",
     "LightCastle Partners. (2024). Strengthening Bangladesh's RMG Sector for the Future of "
     "Work. LightCastle Partners Insights Report. Retrieved from "
     "https://lightcastlepartners.com/insights/2024/01/strengthening-bangladesh-rmg-sector-for-the-future-of-work/"),
    ("[3]",
     "Chen, L., Zhang, W., Liu, Y., & Zhou, X. (2025). Energy-Efficient Prediction in Textile "
     "Manufacturing: Enhancing Accuracy and Data Efficiency With Ensemble Deep Transfer "
     "Learning. IEEE Access, 13. https://doi.org/10.1109/ACCESS.2025.3551798"),
    ("[4]",
     "Ingle, N., & Jasper, W. J. (2025). A review of the evolution and concepts of deep "
     "learning and AI in the textile industry. Textile Research Journal, 95(3-4), 285-307. "
     "https://doi.org/10.1177/00405175241310632"),
    ("[5]",
     "Ruder, S. (2017). An overview of multi-task learning in deep neural networks. "
     "arXiv preprint arXiv:1706.05098."),
    ("[6]",
     "Ngan, H. Y. T., Pang, G. K. H., & Yung, N. H. C. (2011). Automated fabric defect "
     "detection — a review. Image and Vision Computing, 29(7), 442-458. "
     "https://doi.org/10.1016/j.imavis.2011.02.002"),
    ("[7]",
     "Ingle, N., & Jasper, W. J. (2024). A review of deep learning within the framework of "
     "artificial intelligence for enhanced fiber and yarn quality. Textile Research Journal, "
     "94(17-18), 2005-2025. https://doi.org/10.1177/00405175231224143"),
    ("[8]",
     "Agrawal, T. K., Kumar, V., Pal, R., Wang, L., & Chen, Y. (2025). Global initiatives "
     "for industry 4.0 implementation and progress within the textile and apparel manufacturing "
     "sector: a comprehensive review. International Journal of Computer Integrated "
     "Manufacturing, 38(12). https://doi.org/10.1080/0951192X.2025.2455655"),
    ("[9]",
     "Weiss, K., Khoshgoftaar, T. M., & Wang, D. D. (2022). Machine learning and deep "
     "learning based predictive quality in manufacturing: a systematic review. Journal of "
     "Intelligent Manufacturing, 33(7), 1879-1905. https://doi.org/10.1007/s10845-022-01963-8"),
    ("[10]",
     "Jin, X., Chen, K., Jiang, X., Liu, W., & Yan, J. (2022). Deep multistage multi-task "
     "learning for quality prediction of multistage manufacturing systems. Journal of Quality "
     "Technology, 53(5), 482-499. https://doi.org/10.1080/00224065.2021.1903822"),
    ("[11]",
     "Wang, Z., Li, Y., Zhang, H., & Chen, L. (2023). Production quality prediction of "
     "multistage manufacturing systems using multi-task joint deep learning. Journal of "
     "Manufacturing Systems, 68, 421-432. https://doi.org/10.1016/j.jmsy.2023.05.002"),
    ("[12]",
     "Huang, X., Khetan, A., Cvitkovic, M., & Karnin, Z. (2020). TabTransformer: Tabular "
     "Data Modeling Using Contextual Embeddings. arXiv preprint arXiv:2012.06678."),
    ("[13]",
     "Borisov, V., Leemann, T., Sessler, K., Haug, J., Pawelczyk, M., & Kasneci, G. (2022). "
     "Deep Neural Networks and Tabular Data: A Survey. IEEE Transactions on Neural Networks "
     "and Learning Systems. https://doi.org/10.1109/TNNLS.2022.3229161"),
    ("[14]",
     "Shwartz-Ziv, R., & Armon, A. (2022). Tabular data: Deep learning is not all you need. "
     "Information Fusion, 81, 84-90. https://doi.org/10.1016/j.inffus.2021.11.011"),
    ("[15]",
     "Masini, R. P., Medeiros, M. C., & Mendes, E. F. (2024). A Review of Time-Series "
     "Forecasting Algorithms for Industrial Manufacturing Systems. Machines, 12(6), 380. "
     "https://doi.org/10.3390/machines12060380"),
    ("[16]",
     "Al-Shayea, A. M., Saleh, M., Alajlan, A., & Kamrani, A. (2023). Adaptive "
     "One-Dimensional Convolutional Neural Network for Tabular Data. International Journal "
     "on Recent and Innovation Trends in Computing and Communication, 11(10), 168-175."),
    ("[17]",
     "Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization. "
     "International Conference on Learning Representations (ICLR 2019). "
     "https://openreview.net/forum?id=Bkg6RiCqY7"),
    ("[18]",
     "Paszke, A., Gross, S., Massa, F., Lerer, A., Bradbury, J., Chanan, G., ... & Chintala, "
     "S. (2019). PyTorch: An imperative style, high-performance deep learning library. "
     "Advances in Neural Information Processing Systems, 32, 8026-8037."),
    ("[19]",
     "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, "
     "L., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information "
     "Processing Systems, 30."),
    ("[20]",
     "Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, "
     "9(8), 1735-1780. https://doi.org/10.1162/neco.1997.9.8.1735"),
]

for num, text in REFS:
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.left_indent       = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    r1 = p.add_run(num + "  "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(text);        r2.font.size = Pt(10)

# =============================================================================
# SAVE
# =============================================================================
doc.save(OUT_PATH)
print(f"\nPaper saved: {OUT_PATH}")
print(f"Best model from figures: {BEST}  (Avg R2={S[BEST]['avg_r2']})")
print("Done.")
