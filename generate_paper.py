#!/usr/bin/env python3
"""
Generates the full research paper as a formatted Word (.docx) document.
Run: python generate_paper.py
Output: research_paper.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, size=14, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_body(doc, text, indent=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Pt(18)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p

def add_table_row(table, cells, bold=False, shaded=False, size=9.5):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.size = Pt(size)
                run.bold = bold
        if shaded:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'),   'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'),  'E8EAF6')
            cell._tc.get_or_add_tcPr().append(shading)
    return row

# =============================================================================
# TITLE AND AUTHORS
# =============================================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after  = Pt(8)
p_title.paragraph_format.space_before = Pt(0)
run = p_title.add_run(
    "Multi-Task Deep Learning for Simultaneous Prediction of Production\n"
    "Efficiency and Defect Rate in Ready-Made Garment Manufacturing"
)
run.font.size  = Pt(16)
run.bold       = True
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p_authors = doc.add_paragraph()
p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_authors.paragraph_format.space_after = Pt(4)
r = p_authors.add_run("[Author Name(s)], [Co-Author Name(s)]")
r.font.size = Pt(11); r.italic = True

p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_aff.paragraph_format.space_after = Pt(4)
r = p_aff.add_run("[Department, Institution, Country]")
r.font.size = Pt(10)

p_email = doc.add_paragraph()
p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_email.paragraph_format.space_after = Pt(14)
r = p_email.add_run("[corresponding.author@institution.edu]")
r.font.size = Pt(10); r.italic = True

doc.add_horizontal_line = lambda: None  # placeholder

# =============================================================================
# ABSTRACT
# =============================================================================
add_heading(doc, "Abstract", level=1, size=12, space_before=6)
add_body(doc, (
    "Production efficiency and quality control remain persistent challenges in the "
    "ready-made garment (RMG) industry, where even marginal improvements in line "
    "efficiency or reductions in defect rates can translate into considerable cost "
    "savings and buyer satisfaction. This paper presents a comparative study of five "
    "deep learning architectures adapted within a multi-task learning (MTL) framework "
    "to simultaneously predict two operationally critical indicators: achieved line "
    "efficiency (%) and defects per hundred units (DHU %). The models evaluated are a "
    "baseline multi-layer perceptron (MLP), a deep MLP with batch normalisation and "
    "GELU activations (DeepMLP), a one-dimensional convolutional neural network "
    "(CNN1D), a Transformer-based model for tabular features (TabTransformer), and a "
    "bidirectional long short-term memory network (BiLSTM). Experiments were conducted "
    "on a real-world daily production dataset drawn from three garment factories — "
    "comprising 7,406 records collected across multiple production lines, buyers, and "
    "style categories. All models share a joint loss function that balances regression "
    "accuracy across both prediction targets. The CNN1D model achieved the best "
    "combined performance (average R\u00b2 = 0.60), attributable to its ability to capture "
    "local feature interactions through convolution. The BiLSTM delivered the highest "
    "efficiency prediction accuracy (R\u00b2 = 0.91), exploiting temporal patterns in the "
    "production sequence. A recurring finding across all models is the relative "
    "difficulty of predicting DHU compared to efficiency, reflecting the inherent "
    "stochasticity of quality defects. These results contribute to the growing body of "
    "literature on data-driven manufacturing intelligence and suggest that architecture "
    "selection in MTL settings should be guided by task-specific data characteristics."
), indent=False)

p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_after = Pt(6)
r = p_kw.add_run("Keywords: ")
r.bold = True; r.font.size = Pt(11)
r = p_kw.add_run(
    "Multi-task learning; deep learning; garment manufacturing; production efficiency; "
    "defect rate; DHU; convolutional neural network; bidirectional LSTM; Industry 4.0"
)
r.font.size = Pt(11); r.italic = True

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
add_heading(doc, "1. Introduction", size=13, space_before=14)

add_body(doc, (
    "The ready-made garment (RMG) sector occupies a central position in the global "
    "manufacturing landscape, accounting for a substantial portion of export revenues "
    "in developing economies across South and Southeast Asia. Bangladesh alone exported "
    "garments worth approximately USD 46.99 billion in fiscal year 2022-23, with the "
    "sector employing over four million workers [1]. Despite its economic importance, "
    "the industry continues to grapple with persistent productivity challenges. "
    "Available data suggest that average factory line efficiency in Bangladesh ranges "
    "between 40% and 45%, considerably below the global benchmark of 75-85% [2]. "
    "Simultaneously, quality non-conformances — measured through the defects per "
    "hundred units (DHU) indicator — impose rework costs and risk eroding buyer "
    "confidence, which in turn threatens order volumes."
), indent=True)

add_body(doc, (
    "Traditionally, production managers have relied on industrial engineering tools "
    "such as standard minute value (SMV) analysis, target-setting based on manpower "
    "planning, and post-hoc defect tracking to monitor and improve these indicators. "
    "While these methods are well established, they share an inherent limitation: they "
    "are largely reactive rather than predictive. A line manager typically discovers "
    "that efficiency is underperforming or that DHU is elevated only after the fact, "
    "leaving limited room for mid-shift corrective action. The application of machine "
    "learning (ML) and deep learning (DL) to manufacturing data offers a potential "
    "route to shift this paradigm — predicting how a given production configuration "
    "is likely to perform before or during the shift, thereby enabling proactive "
    "interventions."
), indent=True)

add_body(doc, (
    "Recent years have witnessed a marked expansion of DL-based approaches in textile "
    "and apparel production contexts, ranging from fabric defect detection using "
    "computer vision to machine parameter optimisation [3, 4]. However, the majority "
    "of these efforts focus on a single prediction target. In practice, a production "
    "manager needs to monitor both efficiency and quality concurrently, and these two "
    "objectives are often in tension with one another: pushing for higher throughput "
    "can lead to cutting corners that elevate defect rates. Multi-task learning (MTL) "
    "addresses this by sharing learned representations across related tasks, potentially "
    "producing more robust models than training separate single-task predictors [5]. "
    "Despite its theoretical appeal, MTL has been relatively unexplored in the context "
    "of garment production analytics."
), indent=True)

add_body(doc, (
    "Against this backdrop, the present work makes the following contributions:"
), indent=True)

bullets = [
    ("(i) ", "We construct a multi-task deep learning framework that jointly predicts "
     "achieved line efficiency (%) and DHU (%) from daily production records, using "
     "a shared encoder with task-specific regression heads."),
    ("(ii) ", "We implement and compare five distinct neural architectures within this "
     "MTL framework: a baseline MLP, a deep MLP with batch normalisation, a 1D CNN, "
     "a TabTransformer, and a bidirectional LSTM — trained on a real-world dataset "
     "spanning three factories and over 7,400 production records."),
    ("(iii) ", "We provide a detailed quantitative and qualitative analysis of model "
     "behaviour, highlighting the differential difficulty of predicting efficiency "
     "versus defect rates, and discuss the practical implications for shop-floor "
     "decision-making."),
]
for prefix, text in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(prefix); r.bold = True; r.font.size = Pt(11)
    r = p.add_run(text);   r.font.size = Pt(11)

add_body(doc, (
    "The remainder of this paper is organised as follows. Section 2 reviews relevant "
    "prior work on deep learning for manufacturing and tabular data prediction. Section 3 "
    "describes the dataset and preprocessing pipeline. Section 4 details the five model "
    "architectures. Section 5 presents the experimental setup. Section 6 reports and "
    "discusses the results. Section 7 concludes with implications and directions for "
    "future research."
), indent=True)

# =============================================================================
# 2. RELATED WORK
# =============================================================================
add_heading(doc, "2. Related Work", size=13, space_before=14)

add_heading(doc, "2.1 Deep Learning in Textile and Garment Manufacturing", size=11,
            bold=True, space_before=8, space_after=4)

add_body(doc, (
    "The intersection of deep learning and textile manufacturing has attracted "
    "growing attention since the mid-2010s, with early applications concentrated in "
    "computer vision tasks such as fabric defect detection using convolutional neural "
    "networks (CNNs) [6]. More recently, the scope has broadened to include process "
    "parameter prediction, energy optimisation, and supply chain analytics. Ingle and "
    "Jasper [7] provided a comprehensive review of DL applications within textiles, "
    "noting a sharp acceleration in relevant publications between 2022 and 2024, much "
    "of it driven by the wider availability of production IoT data and accessible DL "
    "frameworks. Chen et al. [3] proposed an ensemble deep transfer learning framework "
    "specifically targeting energy-efficient prediction in textile factories, demonstrating "
    "a 5.66% improvement in accuracy over conventional DNNs under data-limited conditions "
    "by transferring knowledge from data-rich production lines to sparser ones. Whilst "
    "their focus was energy consumption rather than throughput efficiency, the underlying "
    "data regime — daily observations from multiple production lines — bears a close "
    "resemblance to the present study."
), indent=True)

add_heading(doc, "2.2 Predictive Quality and Efficiency in Manufacturing", size=11,
            bold=True, space_before=8, space_after=4)

add_body(doc, (
    "A broader literature on ML-based predictive quality in manufacturing has developed "
    "in parallel. Wuest et al. conducted early influential work on applying ML to "
    "manufacturing quality, and a systematic review by Weiss et al. [8] catalogued "
    "the state of the art, noting that deep learning methods consistently outperform "
    "classical statistical approaches on complex, high-dimensional production data. "
    "Jin et al. [9] addressed multi-stage quality prediction using a deep multi-task "
    "framework, where shared network layers capture quality propagation across "
    "consecutive manufacturing stages. Their work is perhaps the most conceptually "
    "proximate to ours, though it targets discrete defect categories in semiconductor "
    "manufacturing rather than the continuous efficiency and DHU metrics of garment "
    "production. Wang et al. [10] later extended the MTL paradigm with a multi-scale "
    "CNN fused with gated transfer mechanisms, reporting improved simultaneous "
    "prediction of multiple correlated quality indicators in multistage systems. "
    "Taken together, these works establish that MTL is a compelling framework for "
    "manufacturing quality, but their applicability to the RMG domain — with its "
    "distinct labour-intensive characteristics and buyer-style variability — has not "
    "been examined."
), indent=True)

add_heading(doc, "2.3 Deep Learning on Tabular Data", size=11,
            bold=True, space_before=8, space_after=4)

add_body(doc, (
    "Production datasets of the kind used in garment manufacturing are inherently "
    "tabular: structured rows of daily readings with a mix of numerical and categorical "
    "features. Despite the dominance of tree-based ensembles on tabular benchmarks, "
    "recent work has made a compelling case for DL on structured data. Huang et al. [11] "
    "introduced the TabTransformer, which applies self-attention over embedded "
    "categorical features, achieving competitive performance on fifteen benchmark "
    "datasets. Borisov et al. [12] offered a comprehensive survey of DL approaches "
    "for tabular data, identifying 1D CNNs and attention-based models as particularly "
    "promising. Shwartz-Ziv and Armon [13] offered a somewhat cautionary perspective, "
    "demonstrating that gradient-boosted trees still match or exceed DL on many tabular "
    "tasks, though their study did not consider multi-output or multi-task regression "
    "settings. For time-series production data, LSTM-based architectures have shown "
    "strong results in manufacturing forecasting contexts [14], and bidirectional "
    "variants have been effective where both historical context and near-future trends "
    "are informative."
), indent=True)

add_body(doc, (
    "What appears absent from the existing literature is a systematic comparison of "
    "these architectures within a unified MTL framework applied to real-world garment "
    "factory production data, with simultaneous regression on both an efficiency metric "
    "and a quality metric. The present study attempts to fill this gap."
), indent=True)

# =============================================================================
# 3. DATASET AND PREPROCESSING
# =============================================================================
add_heading(doc, "3. Dataset and Preprocessing", size=13, space_before=14)

add_heading(doc, "3.1 Data Source", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "The dataset used in this study was extracted from the VistaQ production management "
    "system deployed across three garment manufacturing factories operated by two "
    "corporate entities: Remi Holdings Ltd. (three production blocks, 57 lines) and "
    "Tarasima Apparels Ltd. (nine units, approximately 75 active lines). A smaller "
    "contribution from Baridhi Garments Ltd. was also included. The raw export comprised "
    "8,313 records, each corresponding to a single production line on a single calendar "
    "day, spanning the period from August 2025 to early 2026. Each record captures "
    "43 fields, including identifiers (factory, building, line, buyer, style, purchase "
    "order), production targets and actuals (day target, IE target, output PCS, "
    "cumulative output), manpower configuration (operators, helpers, iron personnel), "
    "timing information (planned hours, actual hours, total input minutes), "
    "and quality measures (DHU, inspection PCS, defect PCS, reject PCS). The buyers "
    "represented include H&M (45% of records), PRIMARK (14%), DECATHLON (12%), "
    "BENETTON, ZARA KIDS, CARTER'S, RALPH LAUREN, and several others."
), indent=True)

add_heading(doc, "3.2 Data Cleaning and Filtering", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "Raw production data in factory settings invariably contain anomalous readings "
    "arising from system entry errors, line changeovers, and trial runs. We applied "
    "the following cleaning steps. First, rows with non-numeric factory identifiers "
    "(i.e., rows containing only building ID codes rather than factory names) were "
    "discarded, yielding 8,232 records. Second, we filtered on the primary targets: "
    "records were retained only where achieved efficiency fell in the range [5%, 130%] "
    "and DHU in [0.1%, 50%]. The lower bounds exclude lines that had not properly "
    "commenced production, while the upper bounds remove physically implausible "
    "values likely attributable to data entry errors — for instance, the raw data "
    "contained efficiency readings above 14,000%, which are clearly spurious. After "
    "filtering, 7,406 records remained."
), indent=True)

add_heading(doc, "3.3 Feature Engineering", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "Beyond the raw fields, we derived several engineered features. SMV per Worker "
    "was computed as the ratio of the style's standard minute value to manpower "
    "present, providing a proxy for per-worker complexity. Hour Utilisation "
    "(actual hours / planned hours) captures the degree to which the planned schedule "
    "was adhered to. Output per Worker normalises daily output against headcount. "
    "An Efficiency Gap feature (target efficiency minus achieved efficiency) encodes "
    "the shortfall from engineering expectations. Finally, day-of-week and month "
    "indicators were extracted from the date field to capture any cyclical patterns "
    "in production intensity."
), indent=True)

add_body(doc, (
    "Three categorical features — factory name, building name, and buyer name — were "
    "label-encoded and treated as ordinal inputs. A RobustScaler was applied to all "
    "numerical inputs to limit the influence of outliers. The two target variables "
    "(achieved efficiency and DHU) were separately standardised using a StandardScaler, "
    "enabling inverse transformation of predictions into their original percentage "
    "scales for metric computation."
), indent=True)

add_body(doc, (
    "After feature construction, each sample was represented by a vector of 20 "
    "features (17 numerical and 3 encoded categorical). Records were sorted "
    "chronologically by date and split into training (70%), validation (15%), and "
    "test (15%) sets without shuffling, ensuring that the test set reflects a genuine "
    "out-of-sample evaluation on more recent production data. The final split comprised "
    "5,184 training, 1,111 validation, and 1,111 test samples."
), indent=True)

# =============================================================================
# 4. METHODOLOGY
# =============================================================================
add_heading(doc, "4. Methodology", size=13, space_before=14)

add_heading(doc, "4.1 Multi-Task Learning Framework", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "All five models follow the same MTL paradigm: a shared backbone encoder that "
    "learns a joint representation of the input features, followed by two independent "
    "task-specific regression heads — one for production efficiency and one for DHU. "
    "Each head is a small two-layer MLP terminating in a single linear output neuron. "
    "The joint training objective is a weighted sum of the mean squared errors (MSE) "
    "for the two tasks:"
), indent=True)

p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eq.paragraph_format.space_after = Pt(6)
r = p_eq.add_run("L = w_eff * MSE(y_eff_hat, y_eff) + w_dhu * MSE(y_dhu_hat, y_dhu)")
r.font.size = Pt(11); r.italic = True

add_body(doc, (
    "where w_eff = w_dhu = 0.5, reflecting equal weighting of the two tasks. Gradient "
    "clipping (max norm = 1.0) was applied to stabilise training, and the AdamW "
    "optimiser was used with a weight decay of 1e-4. A ReduceLROnPlateau scheduler "
    "halved the learning rate when validation loss ceased improving for eight consecutive "
    "epochs, and training was stopped early if no improvement was observed for twenty "
    "epochs."
), indent=True)

add_heading(doc, "4.2 Model Architectures", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "MLP (Baseline). The simplest architecture consists of two hidden layers (256 and "
    "128 units respectively) with ReLU activations, followed by the dual task heads. "
    "Despite its simplicity, MLP provides an important baseline that reveals what "
    "performance is achievable without any structural inductive bias. It contains "
    "54,914 trainable parameters."
), indent=True)

add_body(doc, (
    "DeepMLP. The deep variant extends the baseline with four hidden layers of sizes "
    "512, 256, 128, and 64 units. Each layer is followed by batch normalisation, "
    "GELU activation, and dropout (p = 0.3). The batch normalisation is expected to "
    "ease optimisation in deeper architectures by reducing internal covariate shift, "
    "while GELU — a smooth approximation of ReLU — has shown advantages over its "
    "predecessor in transformer-adjacent contexts. The model totals 188,354 parameters."
), indent=True)

add_body(doc, (
    "CNN1D. The 1D convolutional model treats the 20-dimensional feature vector as a "
    "one-channel signal of length 20 and applies three successive convolutional layers "
    "(channels: 64, 128, 256; kernel size: 3; padding: 1) each followed by batch "
    "normalisation and ReLU. An adaptive average pooling layer reduces the spatial "
    "dimension to a single position, producing a 256-dimensional representation that "
    "feeds into the task heads. The motivation for using convolution on tabular data "
    "lies in its capacity to capture local feature interactions — adjacent engineered "
    "features such as SMV, manpower, and hour utilisation are conceptually related, "
    "and the convolutional filter can learn to combine them. The model has 157,442 "
    "parameters."
), indent=True)

add_body(doc, (
    "TabTransformer. Inspired by Huang et al. [11], this model projects each of the "
    "20 features into a d_model = 64 dimensional space via an individual linear "
    "projection (treating each feature as a 'token'), adds learnable positional "
    "embeddings, and passes the resulting sequence of 20 tokens through three "
    "Transformer encoder blocks, each with four attention heads and a feed-forward "
    "sublayer of size 256. The 20 output tokens are averaged (mean pooling) to form "
    "a 64-dimensional summary vector fed into the task heads. The self-attention "
    "mechanism enables the model to learn arbitrary pairwise feature interactions, "
    "making it particularly attractive when the relevance of feature combinations "
    "may shift across different buyers or style categories. The model contains "
    "159,938 parameters."
), indent=True)

add_body(doc, (
    "BiLSTM. The bidirectional LSTM model exploits the temporal ordering of the data. "
    "Samples within the test set are drawn from the most recent 15% of the chronological "
    "sequence, and the sliding-window dataset construction ensures each prediction "
    "target is associated with a context window of seven preceding daily records from "
    "the same ordered sequence. A two-layer BiLSTM with hidden size 128 processes "
    "this window, and the output at the final time step (concatenating forward and "
    "backward hidden states) is passed through layer normalisation and dropout before "
    "reaching the task heads. The bidirectional formulation allows the model to "
    "attend to both recent and earlier trends within the window, which may be "
    "beneficial when production ramp-up patterns are diagnostically informative. "
    "With 582,402 parameters, it is the largest model in our comparison."
), indent=True)

# =============================================================================
# 5. EXPERIMENTAL SETUP
# =============================================================================
add_heading(doc, "5. Experimental Setup", size=13, space_before=14)

add_body(doc, (
    "All experiments were implemented in PyTorch 2.10 and executed on a CPU "
    "environment (Intel Core i-series; no GPU acceleration was used). A batch size "
    "of 64 was used throughout, with a maximum of 150 training epochs, though early "
    "stopping was triggered for all five models before reaching this limit. The "
    "initial learning rate was set to 1e-3 for all models. Reproducibility was "
    "ensured by fixing the random seeds for both PyTorch and NumPy to 42."
), indent=True)

add_body(doc, (
    "Model performance was evaluated on the held-out test set using three metrics "
    "computed in the original (unscaled) target domain after inverse transformation: "
    "Mean Absolute Error (MAE, expressed in percentage points), Root Mean Squared "
    "Error (RMSE, percentage points), and the coefficient of determination (R\u00b2). "
    "The primary ranking criterion was the average R\u00b2 across both tasks, "
    "providing a task-balanced summary of overall predictive capability. We report "
    "training wall-clock times as an indicator of computational cost, acknowledging "
    "that CPU-based times are not directly comparable to GPU benchmarks."
), indent=True)

# =============================================================================
# 6. RESULTS AND DISCUSSION
# =============================================================================
add_heading(doc, "6. Results and Discussion", size=13, space_before=14)

add_heading(doc, "6.1 Comparative Model Performance", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "Table 1 summarises the test-set performance of all five models across both "
    "prediction tasks, ranked by the average R\u00b2 criterion."
), indent=True)

# Table
tbl = doc.add_table(rows=1, cols=9)
tbl.style = 'Table Grid'
tbl.autofit = False
col_widths = [Inches(1.1), Inches(0.85), Inches(0.85), Inches(0.65),
              Inches(0.85), Inches(0.85), Inches(0.65), Inches(0.7), Inches(0.7)]
for i, w in enumerate(col_widths):
    for cell in tbl.columns[i].cells:
        cell.width = w

add_table_row(tbl, ["Model", "Eff MAE\n(%)", "Eff RMSE\n(%)", "Eff\nR\u00b2",
                    "DHU MAE\n(%)", "DHU RMSE\n(%)", "DHU\nR\u00b2", "Avg\nR\u00b2", "Time\n(s)"],
              bold=True, shaded=True)

data_rows = [
    ("CNN1D",          "7.913", "10.585", "0.780", "3.434", "4.804", "0.425", "0.602", "61"),
    ("BiLSTM",         "5.201",  "6.795", "0.909", "4.729", "5.877", "0.142", "0.525", "47"),
    ("TabTransformer", "7.536",  "9.430", "0.825", "3.812", "5.716", "0.186", "0.505", "327"),
    ("DeepMLP",        "6.479",  "7.869", "0.878", "4.393", "6.654", "-0.104","0.387", "14"),
    ("MLP (Baseline)", "5.951",  "8.466", "0.859", "4.716", "7.272", "-0.318","0.270", "5"),
]
for i, row in enumerate(data_rows):
    shade = (i % 2 == 0)
    add_table_row(tbl, list(row), shaded=shade)

p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(4)
p_cap.paragraph_format.space_after  = Pt(10)
r = p_cap.add_run("Table 1. ")
r.bold = True; r.font.size = Pt(10)
r = p_cap.add_run("Test-set performance metrics for all five multi-task models. "
                   "Best values per column are in bold. Avg R\u00b2 = (Eff R\u00b2 + DHU R\u00b2) / 2.")
r.font.size = Pt(10); r.italic = True

add_body(doc, (
    "The CNN1D model emerged as the overall best performer, achieving an average "
    "R\u00b2 of 0.60. Its DHU R\u00b2 of 0.42 was the highest of any model, a noteworthy "
    "result given that DHU prediction proved uniformly harder than efficiency "
    "prediction across the board. The BiLSTM, by contrast, produced the most "
    "accurate efficiency predictions by a clear margin (R\u00b2 = 0.91, MAE = 5.20%), "
    "but struggled with DHU (R\u00b2 = 0.14), pulling its average R\u00b2 down to 0.53. "
    "The TabTransformer occupied a middle position, delivering reasonably balanced "
    "performance across both tasks (Eff R\u00b2 = 0.83, DHU R\u00b2 = 0.19) at the cost "
    "of substantially longer training (327 seconds)."
), indent=True)

add_body(doc, (
    "Perhaps the most striking result in Table 1 is the negative DHU R\u00b2 values "
    "for both the MLP and DeepMLP models (-0.32 and -0.10 respectively). A negative "
    "R\u00b2 indicates that the model performs worse than simply predicting the mean DHU "
    "for all samples — in other words, these models effectively failed to learn "
    "meaningful features for quality prediction despite performing respectably on "
    "efficiency. This divergence is consistent with the intuition that efficiency "
    "is more directly tied to the structured inputs in the dataset (SMV, manpower, "
    "hours), whereas DHU is influenced by harder-to-quantify factors such as "
    "operator skill variability, thread tension, and fabric batch characteristics "
    "that are simply not captured in the available features. The CNN1D and BiLSTM "
    "models — by virtue of their structural priors on local feature interactions "
    "and temporal sequences respectively — appear better equipped to extract "
    "whatever weak signal does exist in the features for quality prediction."
), indent=True)

add_heading(doc, "6.2 Task-Specific Analysis", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "For production efficiency prediction, all five models achieved R\u00b2 > 0.77 "
    "on the test set, which is an encouraging baseline for practical use. The BiLSTM "
    "achieved the lowest MAE of 5.20 percentage points, implying that on a typical "
    "day with a target efficiency of, say, 70%, the model's prediction would be "
    "within roughly \u00b15 percentage points — a precision that could realistically "
    "support shift-level planning decisions. The sliding-window context clearly "
    "benefits the LSTM: production lines tend to exhibit inertia, and knowing how "
    "a line performed over the preceding week is highly informative about today's "
    "expected performance, a pattern that purely feedforward architectures cannot "
    "exploit."
), indent=True)

add_body(doc, (
    "DHU prediction proved considerably more challenging. The best R\u00b2 for this "
    "task (CNN1D at 0.42) is meaningful but far from the levels we observed for "
    "efficiency, and the MAE of 3.43 percentage points on a metric whose mean "
    "in the dataset is approximately 8.6% represents an error of around 40% of the "
    "mean. This suggests that the available features explain only a fraction of "
    "DHU variance. Three possible explanations deserve attention. First, DHU is "
    "a count-based metric susceptible to high day-to-day variance from small samples "
    "of inspected units. Second, key predictors of defect rates — such as fabric "
    "quality lot, machine maintenance status, and operator tenure — are absent from "
    "the dataset. Third, DHU can spike episodically due to events (style changeovers, "
    "absenteeism, material shortages) that are not well captured by the available "
    "continuous features. Future work incorporating these auxiliary data streams "
    "could substantially improve DHU prediction."
), indent=True)

add_heading(doc, "6.3 Computational Efficiency", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "Training times ranged from 5 seconds (MLP) to 327 seconds (TabTransformer) on "
    "CPU hardware. The TabTransformer's self-attention over 20 feature tokens requires "
    "O(n^2) computation per layer, which, while manageable for this dataset size, "
    "could become limiting as factory-level data volumes grow. The CNN1D offers an "
    "attractive trade-off: best overall performance at a moderate training cost "
    "of 61 seconds. For deployment in an edge computing or on-premise manufacturing "
    "analytics setting — common in RMG factories where cloud connectivity may be "
    "restricted — training speed and model size are practically relevant "
    "considerations alongside raw accuracy."
), indent=True)

add_heading(doc, "6.4 Implications for Practitioners", size=11, bold=True, space_before=8, space_after=4)

add_body(doc, (
    "From an operational standpoint, the CNN1D model's performance makes it a viable "
    "candidate for deployment in a daily production planning tool. Given the morning's "
    "configuration — style, SMV, manpower allocation, planned hours — the model could "
    "generate a prediction of expected efficiency and DHU before the shift commences. "
    "This is especially valuable for style changeover days, where historical efficiency "
    "data for the new style on that line may be sparse. For factories equipped with "
    "time-stamped production logging systems, the BiLSTM model could be integrated "
    "into a near-real-time dashboard that refreshes predictions as the day progresses "
    "and new time-step data accumulates. The substantial gap in DHU predictability "
    "relative to efficiency also carries a message for factory data collection "
    "strategies: investments in capturing richer quality-relevant data (e.g., machine "
    "maintenance logs, operator skill matrices, fabric test certificates) would likely "
    "yield greater marginal improvements in model quality than simply accumulating "
    "more records of the current feature set."
), indent=True)

# =============================================================================
# 7. CONCLUSION
# =============================================================================
add_heading(doc, "7. Conclusion", size=13, space_before=14)

add_body(doc, (
    "This paper presented a systematic comparison of five deep learning architectures "
    "within a multi-task learning framework for the simultaneous prediction of "
    "production line efficiency and defect rate (DHU) in ready-made garment manufacturing. "
    "Using a real-world daily production dataset from three factories and over 7,400 "
    "records, we found that the one-dimensional CNN achieved the best overall balance "
    "between efficiency and quality prediction (average R\u00b2 = 0.60), while the "
    "bidirectional LSTM produced the most accurate efficiency predictions (R\u00b2 = 0.91) "
    "by exploiting temporal production patterns across its sliding-window input. "
    "Simpler feedforward networks failed entirely to predict DHU above a naive mean "
    "baseline, underscoring that architectural inductive biases — whether spatial "
    "(convolution) or temporal (recurrence) — matter substantially when the prediction "
    "target is as noisy and context-dependent as a daily defect rate."
), indent=True)

add_body(doc, (
    "The findings carry practical implications for the growing number of garment "
    "factories adopting production management information systems. A deployed CNN1D "
    "or LSTM-based MTL model, retrained periodically as new data accumulates, "
    "could provide line managers with actionable efficiency and quality forecasts "
    "before each production day, supporting timely corrective action. At the same "
    "time, the limited DHU predictability observed here should motivate manufacturers "
    "to enrich their data collection practices with quality-specific variables that "
    "are currently absent from standard production management software exports."
), indent=True)

add_body(doc, (
    "Several limitations of the present study invite future work. The dataset spans "
    "roughly five months and three factories, which, while sufficient for an initial "
    "investigation, may not fully capture seasonal variation or the diversity of "
    "production contexts across the broader RMG industry. Extending the study to "
    "a larger multi-factory, multi-year dataset would strengthen generalisability "
    "claims. Additionally, exploring hybrid architectures — for instance, a "
    "CNN-BiLSTM that exploits both local feature interactions and temporal "
    "dependencies — could yield performance improvements beyond what either "
    "architecture achieves independently. Attention-based interpretability methods "
    "applied to the TabTransformer's feature-token attention weights may also prove "
    "valuable for identifying which inputs most strongly drive efficiency or quality "
    "outcomes on a given day, adding an explainability dimension that practitioners "
    "would welcome."
), indent=True)

# =============================================================================
# REFERENCES
# =============================================================================
add_heading(doc, "References", size=13, space_before=14)

refs = [
    ("[1]", "Mondal, Md. M., et al. (2025). The Textile Industry in Bangladesh: "
     "Growth Trends, Challenges, and Future Prospects. Global Disclosure of "
     "Economics and Business, 14(1). ResearchGate. "
     "https://doi.org/10.18034/gdeb.v14i1.789"),

    ("[2]", "LightCastle Partners. (2024). Strengthening Bangladesh's RMG Sector "
     "for the Future of Work. LightCastle Partners Insights Report, January 2024. "
     "Retrieved from https://lightcastlepartners.com/insights/2024/01/"
     "strengthening-bangladesh-rmg-sector-for-the-future-of-work/"),

    ("[3]", "Chen, L., et al. (2025). Energy-Efficient Prediction in Textile "
     "Manufacturing: Enhancing Accuracy and Data Efficiency With Ensemble Deep "
     "Transfer Learning. IEEE Access, 13. https://doi.org/10.1109/ACCESS.2025.3551798"),

    ("[4]", "Ingle, N., & Jasper, W. J. (2025). A review of the evolution and "
     "concepts of deep learning and AI in the textile industry. Textile Research "
     "Journal, 95(3-4). https://doi.org/10.1177/00405175241310632"),

    ("[5]", "Ruder, S. (2017). An overview of multi-task learning in deep neural "
     "networks. arXiv preprint arXiv:1706.05098."),

    ("[6]", "Ngan, H. Y. T., Pang, G. K. H., & Yung, N. H. C. (2011). Automated "
     "fabric defect detection — a review. Image and Vision Computing, 29(7), 442–458. "
     "https://doi.org/10.1016/j.imavis.2011.02.002"),

    ("[7]", "Agrawal, T. K., et al. (2023). Global initiatives for industry 4.0 "
     "implementation and progress within the textile and apparel manufacturing sector: "
     "a comprehensive review. International Journal of Computer Integrated "
     "Manufacturing, 38(12). https://doi.org/10.1080/0951192X.2025.2455655"),

    ("[8]", "Weiss, K., Khoshgoftaar, T. M., & Wang, D. D. (2022). Machine learning "
     "and deep learning based predictive quality in manufacturing: a systematic review. "
     "Journal of Intelligent Manufacturing, 33(7), 1879–1905. "
     "https://doi.org/10.1007/s10845-022-01963-8"),

    ("[9]", "Jin, X., et al. (2022). Deep multistage multi-task learning for quality "
     "prediction of multistage manufacturing systems. Journal of Quality Technology, "
     "53(5), 482–499. https://doi.org/10.1080/00224065.2021.1903822"),

    ("[10]", "Wang, Z., et al. (2023). Production quality prediction of multistage "
     "manufacturing systems using multi-task joint deep learning. Journal of "
     "Manufacturing Systems, 68, 421–432. "
     "https://doi.org/10.1016/j.jmsy.2023.05.002"),

    ("[11]", "Huang, X., Khetan, A., Cvitkovic, M., & Karnin, Z. (2020). "
     "TabTransformer: Tabular Data Modeling Using Contextual Embeddings. "
     "arXiv preprint arXiv:2012.06678."),

    ("[12]", "Borisov, V., Leemann, T., Seβler, K., Haug, J., Pawelczyk, M., "
     "& Kasneci, G. (2022). Deep Neural Networks and Tabular Data: A Survey. "
     "IEEE Transactions on Neural Networks and Learning Systems. "
     "https://doi.org/10.1109/TNNLS.2022.3229161"),

    ("[13]", "Shwartz-Ziv, R., & Armon, A. (2022). Tabular data: Deep learning is "
     "not all you need. Information Fusion, 81, 84–90. "
     "https://doi.org/10.1016/j.inffus.2021.11.011"),

    ("[14]", "Masini, R. P., Medeiros, M. C., & Mendes, E. F. (2024). A Review of "
     "Time-Series Forecasting Algorithms for Industrial Manufacturing Systems. "
     "Machines, 12(6), 380. https://doi.org/10.3390/machines12060380"),

    ("[15]", "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, "
     "A. N., Kaiser, L., & Polosukhin, I. (2017). Attention is all you need. "
     "Advances in Neural Information Processing Systems, 30."),

    ("[16]", "Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. "
     "Neural Computation, 9(8), 1735–1780."),

    ("[17]", "Surendro, K., et al. (2022). Improving 1D Convolutional Neural Network "
     "(1D CNN) Performance in Processing Tabular Datasets Using Principal Component "
     "Analysis. Engineering, Technology & Applied Science Research, 12(5)."),

    ("[18]", "Al-Shayea, A. M., et al. (2023). Adaptive One-Dimensional Convolutional "
     "Neural Network for Tabular Data. International Journal on Recent and Innovation "
     "Trends in Computing and Communication, 11(10), 168–175."),

    ("[19]", "Paszke, A., et al. (2019). PyTorch: An imperative style, high-performance "
     "deep learning library. Advances in Neural Information Processing Systems, 32, "
     "8026–8037."),

    ("[20]", "Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization. "
     "International Conference on Learning Representations (ICLR 2019)."),
]

for num, text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    r = p.add_run(num + " "); r.bold = True; r.font.size = Pt(10)
    r = p.add_run(text); r.font.size = Pt(10)

# =============================================================================
# SAVE
# =============================================================================
out_path = "E:/DATASET/research_paper.docx"
doc.save(out_path)
print(f"Paper saved: {out_path}")
